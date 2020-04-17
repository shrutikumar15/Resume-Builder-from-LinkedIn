from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from docx.shared import RGBColor
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.enum.text import WD_COLOR_INDEX
#from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx2pdf import convert
import json

def output1(file_data):
    #file_header = open("scraped_data.json", 'r')

    #file_data = file_header.read()
    json_data = json.loads(file_data)

    basic_info_list = json_data['basic_info_list']
    education_info_list = json_data['education_info_list']
    projects_info_list = json_data['projects_info_list']
    certifications_info_list = json_data['certifications_info_list']
    experience_info_list = json_data['experience_info_list']
    skills_info_list = json_data['skills_info_list']
    volunteer_info_list = json_data['volunteer_info_list']
    accomplishments_info_list = json_data['accomplishments_info_list']
    hobbies_info_list = json_data['hobbies_info_list']

    #print("LISTS")
    #print(basic_info_list)
    #print(education_info_list)
    #print(projects_info_list)
    #print(certifications_info_list)
    #print(experience_info_list)
    #print(skills_info_list)
    #print(volunteer_info_list)
    #print(accomplishments_info_list)
    #ALL LISTS


    document = Document('input1.docx')
    section = document.sections[0]

    #Styles
    ##NameStyle
    style_name = document.styles['Normal']
    font_name = style_name.font
    font_name.size = Pt(14)
    font_name.name = 'Times New Roman'
    style_name.paragraph_format.space_after = Pt(2)
    font_name.bold=True
    ##HeadlineStyle
    styles=document.styles
    #style_headline = styles.add_style('headLine1', WD_STYLE_TYPE.PARAGRAPH)
    style_headline = document.styles['headLine1']
    paragraph_format_headline = style_headline.paragraph_format
    style_headline.paragraph_format.space_after = Pt(2)
    font_headline = style_headline.font
    font_headline.size = Pt(12)
    font_headline.name = 'Times New Roman'
    font_headline.bold=True
    ##EmailAndLinksStyle
    style_email = document.styles['No Spacing']
    style_email.font.size = Pt(9)
    ##LineStyle
    style_line = document.styles['Line1']
    paragraph_format_line = style_line.paragraph_format
    style_line.paragraph_format.space_after = Pt(2)
    font_line = style_line.font
    font_line.size = Pt(1)
    font_line.highlight_color = WD_COLOR_INDEX.BLACK
    font_line.bold=True
    #font_line.color.rgb = RGBColor(31, 73, 125)
    ##EmptySpaceStyle
    style_empty = document.styles['Subtitle']
    font_empty = style_empty.font
    font_empty.size=Pt(1)



    #Function
    ##TABLEINDENTATION
    def indent_table(table, indent):
        # noinspection PyProtectedMember
        tbl_pr = table._element.xpath('w:tblPr')
        if tbl_pr:
            e = OxmlElement('w:tblInd')
            e.set(qn('w:w'), str(indent))
            e.set(qn('w:type'), 'dxa')
            tbl_pr[0].append(e)
    ##TABLEWIDTH
    def ChangeWidthOfTable(table,width,column):
        for columnVarible in range(0,column):
            for cell in table.columns[columnVarible].cells:
                cell.width = Pt(width)
    ###ACADEMICSWIDTH
    def ChangeWidthOfTableAcademics(table,column):
        for columnVariable in range(0,column):
            if columnVariable==0:
                for cell in table.columns[columnVariable].cells:
                    cell.width = Pt(240)
            elif columnVariable==1:
                for cell in table.columns[columnVariable].cells:
                    cell.width = Pt(120)
            else:
                for cell in table.columns[columnVariable].cells:
                    cell.width = Pt(60)
    ###PROJECTSWIDTH
    def ChangeWidthOfTableProjects(table,column):
        for columnVariable in range(0,column):
            if columnVariable==0:
                for cell in table.columns[columnVariable].cells:
                    cell.width = Pt(370)
            elif columnVariable==1:
                for cell in table.columns[columnVariable].cells:
                    cell.width = Pt(110)
    ###CERTIFICATIONSWIDTH
    def ChangeWidthOfTableCertifications(table,column):
        for columnVariable in range(0,column):
            if columnVariable==0:
                for cell in table.columns[columnVariable].cells:
                    cell.width = Pt(400)
            elif columnVariable==1:
                for cell in table.columns[columnVariable].cells:
                    cell.width = Pt(80)
    ##ALLSTYLES
    dstyles=document.styles
    def allstyles(dstyles):
        k = len(dstyles)
        #print (k) 
        for i in dstyles:
            pass
            #print (i.name) ##to print all styles
    allstyles(dstyles)
    ##TABLESTYLE
    def tablesty(table):
        table.rows[0].cells[0]._tc.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="002366"/>'.format(nsdecls('w'))))
        table.rows[0].cells[0].paragraphs[0].runs[0].font.size = Pt(9)
        table.rows[0].cells[0].paragraphs[0].runs[0].font.name = 'Tahoma'
        table.rows[0].cells[0].paragraphs[0].runs[0].font.bold=True
    ##ACADEMICSMAINTABLESTYLE
    def academics_maintablesty(table):
        i=0
        for rows in table.rows:
            for cells in rows.cells:
                cells.paragraphs[0].runs[0].font.size=Pt(10)
                if i==0:
                    cells.paragraphs[0].runs[0].font.bold=True
                else:
                    cells.paragraphs[0].runs[0].font.bold=False
            i+=1
    ##PROJECTMAINTABLESTYLE
    def projects_maintablesty(table):
        i=0
        for rows in table.rows:
            for cells in rows.cells:
                j=0
                if i%2==0:
                    cells.paragraphs[0].runs[0].font.size=Pt(10)
                    cells.paragraphs[0].runs[0].font.bold=True
                elif i%2!=0:
                    if cells.text=="":
                        pass
                    else:
                        cells.paragraphs[0].runs[0].font.size=Pt(10)
                        cells.paragraphs[0].runs[0].font.bold=False
                j+=1
            i+=1
    ##ACCOMPLISHMENTSMAINTABLESTYLE
    def accomplishments_maintablesty(table):
        i=0
        for rows in table.rows:
            for cells in rows.cells:
                j=0
                if i==0:
                    cells.paragraphs[0].runs[0].font.size=Pt(10)
                    cells.paragraphs[0].runs[0].font.bold=True
                else:
                    cells.paragraphs[0].runs[0].font.size=Pt(10)
                    cells.paragraphs[0].runs[0].font.bold=False
                j+=1
            i+=1
    ##CERTIFICATIONSMAINTABLESTYLE
    def certifications_maintablesty(table):
        i=0
        for rows in table.rows:
            for cells in rows.cells:
                cells.paragraphs[0].runs[0].font.size=Pt(10)
                cells.paragraphs[0].runs[0].font.bold=False
            i+=1

    #HEADER
    header = section.header

    sections = document.sections
    ##NAME
    name = header.paragraphs[0]
    name.text = (basic_info_list[0].upper()," ",basic_info_list[1].upper())
    name.paragraph_format.first_line_indent = Pt(-50)
    name.style = document.styles['Normal']
    ##HEADLINE
    if basic_info_list[2]=='':
        pass
    else :
        headLine = header.add_paragraph(basic_info_list[2])
        headLine.paragraph_format.first_line_indent = Pt(-50)
        headLine.style = document.styles['headLine1']
    ##EMAIL
    flag=0
    if basic_info_list[4]=='' and basic_info_list[5]=='' and basic_info_list[6]=='':
        flag=1
    if basic_info_list[4]=='' :
        str1=""
    else:
        str1=basic_info_list[4]+" | "
    if basic_info_list[5]=='' :
        str2=""
    else:
        str2=basic_info_list[5]+" | "
    if basic_info_list[6]=='':
        str3=""
    else:
        str3="Age : "+basic_info_list[6]
    if flag==1:
        pass
    else:
        email = header.add_paragraph(str1+str2+str3)
        email.paragraph_format.first_line_indent = Pt(-50)
        email.style = document.styles['No Spacing']
    ##LINKEDIN
    flag=0
    if basic_info_list[3]=="" and basic_info_list[7]=="":
        flag=1
    if basic_info_list[3]=="":
        str1=""
    else:
        str1="LinkedIn - "+basic_info_list[3]+" | "
    if basic_info_list[7]=="":
        str2=""
    else:
        str2="GitHub - "+basic_info_list[7]
    if flag==1:
        pass
    else:
        linkedin = header.add_paragraph(str1+" "+str2)
        linkedin.paragraph_format.first_line_indent = Pt(-50)
        linkedin.style = document.styles['No Spacing']

    #LINE
    line = document.add_paragraph("_______________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________")
    line.style=document.styles['Line1']
    line.paragraph_format.first_line_indent = Pt(-50)
    #EMPTYSPACE
    es1 = document.add_paragraph()
    es1.paragraph_format.first_line_indent = Pt(-50)
    es1.style = document.styles['Subtitle']
    #TABLE
    ##ACADEMICS
    if len(education_info_list)==0:
        pass
    else:
        #ACADEMICSHEADER
        tabler = document.add_table(rows=1,cols=1)
        tabler.rows[0].cells[0].text="ACADEMICS"
        tabler.style = document.styles['Light List']
        indent_table(tabler,-900)
        ChangeWidthOfTable(tabler,480,1)
        tablesty(tabler)
        ##EMPTYSPACE
        es2 = document.add_paragraph()
        es2.paragraph_format.first_line_indent = Pt(-50)
        es2.style = document.styles['Subtitle']
        ##ACADEMICSTABLE
        row_num = len(education_info_list)
        tablerr = document.add_table(rows=row_num+1,cols=4)
        indent_table(tablerr,-900)
        ChangeWidthOfTableAcademics(tablerr,4)
        tablerr.style = document.styles['Main Table']
        tablerr.rows[0].cells[0].text="Institute"
        tablerr.rows[0].cells[1].text="Board / University"
        tablerr.rows[0].cells[2].text="Year"
        tablerr.rows[0].cells[3].text="% / CGPA"
        for i in range (1,len(tablerr.rows)):
            j=0
            for cells in tablerr.rows[i].cells:
                s = education_info_list[i-1]
                s1 = s[j]
                cells.text=s1
                j+=1
        academics_maintablesty(tablerr)
        ##EMPTYSPACE
        es3 = document.add_paragraph()
        es3.paragraph_format.first_line_indent = Pt(-50)
        es3.style = document.styles['Subtitle']
    ##PROJECTS
    if len(projects_info_list)==0:
        pass
    else:
        ##PROJECTSHEADER
        table_project_header = document.add_table(rows=1,cols=1)
        table_project_header.rows[0].cells[0].text="PROJECTS"
        table_project_header.style = document.styles['Light List']
        indent_table(table_project_header,-900)
        ChangeWidthOfTable(table_project_header,480,1)
        tablesty(table_project_header)
        ##EMPTYSPACE
        es4 = document.add_paragraph()
        es4.paragraph_format.first_line_indent = Pt(-50)
        es4.style = document.styles['Subtitle']
        ##PROJECTSTABLE
        row_num = len(projects_info_list)
        table_project = document.add_table(rows=(row_num)*2,cols=2)
        indent_table(table_project,-900)
        ChangeWidthOfTableProjects(table_project,2)
        table_project.style = document.styles['Main Project Table']
        for i in range (0,len(table_project.rows)):
            j=0
            for cells in table_project.rows[i].cells:
                if i%2==0 and j==0:
                    s = projects_info_list[(int)(i/2)]
                    s1 = "• "+s[j]
                    cells.text=s1
                elif i%2==0 and j==1:
                    s = projects_info_list[(int)(i/2)]
                    s1 = s[j]
                    cells.text=s1
                elif i%2!=0 and j==0 :
                    s=projects_info_list[(int)(i/2)]
                    s1 = s[2]
                    cells.text=s1
                elif i%2!=0 and j==1 :
                    cells.text=''
                j+=1
        projects_maintablesty(table_project)
        ##EMPTYSPACE
        es4 = document.add_paragraph()
        es4.paragraph_format.first_line_indent = Pt(-50)
        es4.style = document.styles['Subtitle']
    ##CERTIFICATIONS
    if len(certifications_info_list)==0:
        pass
    else:
        ##CERTIFICATIONSTSHEADER
        table_certifications_header = document.add_table(rows=1,cols=1)
        table_certifications_header.rows[0].cells[0].text="CERTIFICATIONS"
        table_certifications_header.style = document.styles['Light List']
        indent_table(table_certifications_header,-900)
        ChangeWidthOfTable(table_certifications_header,480,1)
        tablesty(table_certifications_header)
        ##EMPTYSPACE
        es5 = document.add_paragraph()
        es5.paragraph_format.first_line_indent = Pt(-50)
        es5.style = document.styles['Subtitle']
        ##CERTIFICATIONSTABLE
        row_num = len(certifications_info_list)
        table_certifications = document.add_table(rows=(row_num),cols=2)
        indent_table(table_certifications,-900)
        ChangeWidthOfTableCertifications(table_certifications,2)
        table_certifications.style = document.styles['Main Project Table']
        for i in range (0,len(table_certifications.rows)):
            j=0
            for cells in table_certifications.rows[i].cells:
                s = certifications_info_list[i]
                if j==0:
                    s1 = "• "+s[j]+" Issued by "+s[2]
                else :
                    s1 = s[j]
                cells.text=s1
                j+=1
        certifications_maintablesty(table_certifications)    
        ##EMPTYSPACE
        es6 = document.add_paragraph()
        es6.paragraph_format.first_line_indent = Pt(-50)
        es6.style = document.styles['Subtitle']
    ##SKILLS
    if len(skills_info_list)==0:
        pass
    else:
        ##CERTIFICATIONSTSHEADER
        table_skills_header = document.add_table(rows=1,cols=1)
        table_skills_header.rows[0].cells[0].text="TECHNICAL SKILLS"
        table_skills_header.style = document.styles['Light List']
        indent_table(table_skills_header,-900)
        ChangeWidthOfTable(table_skills_header,480,1)
        tablesty(table_skills_header)
        ##EMPTYSPACE
        es7 = document.add_paragraph()
        es7.paragraph_format.first_line_indent = Pt(-50)
        es7.style = document.styles['Subtitle']
        ##SKILLSTABLE
        table_skills = document.add_table(rows=1,cols=1)
        indent_table(table_skills,-900)
        ChangeWidthOfTable(table_skills,480,1)
        table_skills.style = document.styles['Main Project Table']
        for i in range (0,len(table_skills.rows)):
            len1 = len(skills_info_list)
            s1=""
            for len1 in range (0, len1-1):
                s1+=skills_info_list[len1]+", "
            s1+=skills_info_list[len1+1]
            table_skills.rows[0].cells[0].text=s1
        certifications_maintablesty(table_skills)
        ##EMPTYSPACE
        es8 = document.add_paragraph()
        es8.paragraph_format.first_line_indent = Pt(-50)
        es8.style = document.styles['Subtitle']
    ##EXPERIENCE
    if len(experience_info_list)==0:
        pass
    else:
        ##EXPERIENCEHEADER
        table_experience_header = document.add_table(rows=1,cols=1)
        table_experience_header.rows[0].cells[0].text="WORK EXPERIENCE"
        table_experience_header.style = document.styles['Light List']
        indent_table(table_experience_header,-900)
        ChangeWidthOfTable(table_experience_header,480,1)
        tablesty(table_experience_header)
        ##EMPTYSPACE
        es9 = document.add_paragraph()
        es9.paragraph_format.first_line_indent = Pt(-50)
        es9.style = document.styles['Subtitle']
        ##EXPERIENCETABLE
        row_num = len(experience_info_list)
        table_experience = document.add_table(rows=(row_num)*2,cols=2)
        indent_table(table_experience,-900)
        ChangeWidthOfTableProjects(table_experience,2)
        table_experience.style = document.styles['Main Project Table']
        for i in range (0,len(table_experience.rows)):
            j=0
            for cells in table_experience.rows[i].cells:
                if i%2==0 and j==0:
                    s = experience_info_list[(int)(i/2)]
                    s1 = "• "+s[1]+" at "+s[0]
                    cells.text=s1
                elif i%2==0 and j==1:
                    s = experience_info_list[(int)(i/2)]
                    s1 = s[2]
                    cells.text=s1
                elif i%2!=0 and j==0 :
                    s = experience_info_list[(int)(i/2)]
                    s1 = s[3]
                    cells.text=s1
                elif i%2!=0 and j==1 :
                    cells.text=''
                j+=1
        projects_maintablesty(table_experience)  
        ##EMPTYSPACE
        es10 = document.add_paragraph()
        es10.paragraph_format.first_line_indent = Pt(-50)
        es10.style = document.styles['Subtitle']
    ##VOLUNTEER
    if len(volunteer_info_list)==0:
        pass
    else:
        ##EXPERIENCEHEADER
        table_volunteer_header = document.add_table(rows=1,cols=1)
        table_volunteer_header.rows[0].cells[0].text="VOLUNTEER EXPERIENCE"
        table_volunteer_header.style = document.styles['Light List']
        indent_table(table_volunteer_header,-900)
        ChangeWidthOfTable(table_volunteer_header,480,1)
        tablesty(table_volunteer_header)
        ##EMPTYSPACE
        es11 = document.add_paragraph()
        es11.paragraph_format.first_line_indent = Pt(-50)
        es11.style = document.styles['Subtitle']
        ##EXPERIENCETABLE
        row_num = len(volunteer_info_list)
        table_vexperience = document.add_table(rows=(row_num)*2,cols=2)
        indent_table(table_vexperience,-900)
        ChangeWidthOfTableProjects(table_vexperience,2)
        table_vexperience.style = document.styles['Main Project Table']
        for i in range (0,len(table_vexperience.rows)):
            j=0
            for cells in table_vexperience.rows[i].cells:
                if i%2==0 and j==0:
                    s = volunteer_info_list[(int)(i/2)]
                    s1 = "• "+s[1]+" at "+s[0]
                    cells.text=s1
                elif i%2==0 and j==1:
                    s = volunteer_info_list[(int)(i/2)]
                    s1 = s[2]
                    cells.text=s1
                elif i%2!=0 and j==0 :
                    s = volunteer_info_list[(int)(i/2)]
                    s1 = s[3]
                    cells.text=s1
                elif i%2!=0 and j==1 :
                    cells.text=''
                j+=1
        projects_maintablesty(table_vexperience)  
        ##EMPTYSPACE
        es12 = document.add_paragraph()
        es12.paragraph_format.first_line_indent = Pt(-50)
        es12.style = document.styles['Subtitle']
    ##COCURRICULAR
    if  len(accomplishments_info_list)==0:
        pass
    else:
        ##ACCOMPLISHMENTSHEADER
        table_cocurricular_header = document.add_table(rows=1,cols=1)
        table_cocurricular_header.rows[0].cells[0].text="ACCOMPLISHMENTS"
        table_cocurricular_header.style = document.styles['Light List']
        indent_table(table_cocurricular_header,-900)
        ChangeWidthOfTable(table_cocurricular_header,480,1)
        tablesty(table_cocurricular_header)
        ##EMPTYSPACE
        es13 = document.add_paragraph()
        es13.paragraph_format.first_line_indent = Pt(-50)
        es13.style = document.styles['Subtitle']
        ##COCURRICULARTABLE
        if len(accomplishments_info_list)!=0:
        ##ACCOMPLISHMENTSTABLE
            row_num = len(accomplishments_info_list)
            table_accomplishments = document.add_table(rows=(row_num),cols=2)
            indent_table(table_accomplishments,-900)
            ChangeWidthOfTableCertifications(table_accomplishments,2)
            table_accomplishments.style = document.styles['Main Project Table']
            for i in range (0,row_num):
                j=0
                for cells in table_accomplishments.rows[i].cells:
                    if j==0:
                        s = accomplishments_info_list[(int)(i)]
                        s1 = "• "+s[0]
                        cells.text=s1
                    elif j==1:
                        s = accomplishments_info_list[(int)(i)]
                        s1 = s[1]
                        cells.text=s1
                    j+=1
            certifications_maintablesty(table_accomplishments)
            ##EMPTYSPACE
            es14 = document.add_paragraph()
            es14.paragraph_format.first_line_indent = Pt(-50)
            es14.style = document.styles['Subtitle']
        
    ##HOBBIES
    if  len(hobbies_info_list)==0:
        pass
    else:
        ##ACCOMPLISHMENTSHEADER
        table_hobbies_header = document.add_table(rows=1,cols=1)
        table_hobbies_header.rows[0].cells[0].text="HOBBIES"
        table_hobbies_header.style = document.styles['Light List']
        indent_table(table_hobbies_header,-900)
        ChangeWidthOfTable(table_hobbies_header,480,1)
        tablesty(table_hobbies_header)
        ##EMPTYSPACE
        es16 = document.add_paragraph()
        es16.paragraph_format.first_line_indent = Pt(-50)
        es16.style = document.styles['Subtitle']
        if len(hobbies_info_list)!=0:
            ##HOBBIESTABLE
            row_num = len(hobbies_info_list)
            table_hobbies = document.add_table(rows=(row_num),cols=1)
            indent_table(table_hobbies,-900)
            ChangeWidthOfTable(table_hobbies,480,1)
            table_hobbies.style = document.styles['Main Project Table']
            for i in range (0,row_num):
                for cells in table_hobbies.rows[i].cells:
                    s = "• "+hobbies_info_list[i]
                    cells.text=s
            certifications_maintablesty(table_hobbies)
            ##EMPTYSPACE
            es17 = document.add_paragraph()
            es17.paragraph_format.first_line_indent = Pt(-50)
            es17.style = document.styles['Subtitle']


    document.save('input.docx')
    convert("input.docx", "output.pdf")